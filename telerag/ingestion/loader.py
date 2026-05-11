from __future__ import annotations

from pathlib import Path

from telerag.types import Document


class DocumentLoader:
    SUPPORTED_SUFFIXES = {".txt", ".pdf", ".docx", ".html", ".htm", ".md"}

    def load_path(self, docs_path: str | Path) -> list[Document]:
        root = Path(docs_path)
        if not root.exists():
            raise FileNotFoundError(f"Document path does not exist: {root}")

        files = [root] if root.is_file() else sorted(path for path in root.rglob("*") if path.suffix.lower() in self.SUPPORTED_SUFFIXES)
        documents: list[Document] = []
        for file_path in files:
            documents.extend(self._load_file(file_path, root if root.is_dir() else file_path.parent))
        return documents

    def _load_file(self, file_path: Path, base_dir: Path) -> list[Document]:
        suffix = file_path.suffix.lower()
        if suffix == ".txt":
            text = self._read_txt(file_path)
            return self._make_documents(file_path, base_dir, [text])
        if suffix == ".pdf":
            return self._read_pdf(file_path, base_dir)
        if suffix == ".docx":
            text = self._read_docx(file_path)
            return self._make_documents(file_path, base_dir, [text])
        if suffix in {".html", ".htm"}:
            text = self._read_html(file_path)
            return self._make_documents(file_path, base_dir, [text])
        if suffix == ".md":
            text = self._read_markdown(file_path)
            return self._make_documents(file_path, base_dir, [text])
        return []

    def _read_txt(self, file_path: Path) -> str:
        for encoding in ("utf-8", "utf-8-sig", "gb18030", "latin-1"):
            try:
                return file_path.read_text(encoding=encoding)
            except UnicodeDecodeError:
                continue
        raise UnicodeDecodeError("unknown", b"", 0, 1, f"Unable to decode text file: {file_path}")

    def _read_pdf(self, file_path: Path, base_dir: Path) -> list[Document]:
        import fitz

        documents: list[Document] = []
        with fitz.open(file_path) as pdf:
            for page_index, page in enumerate(pdf):
                text = page.get_text("text")
                cleaned = self._clean_text(text)
                if cleaned:
                    documents.append(
                        Document(
                            text=cleaned,
                            metadata={
                                "source": str(file_path.relative_to(base_dir)),
                                "page": page_index + 1,
                            },
                        )
                    )
        return documents

    def _read_docx(self, file_path: Path) -> str:
        from docx import Document as DocxDocument

        document = DocxDocument(file_path)
        paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs]
        return "\n\n".join(text for text in paragraphs if text)

    def _read_html(self, file_path: Path) -> str:
        from bs4 import BeautifulSoup

        html = self._read_txt(file_path)
        soup = BeautifulSoup(html, "lxml")
        for tag in soup(["script", "style", "noscript", "svg", "footer", "nav", "aside", "form"]):
            tag.decompose()

        root = soup.find("main") or soup.find("article") or soup.body or soup
        return self._html_to_text(root)

    def _read_markdown(self, file_path: Path) -> str:
        from markdown_it import MarkdownIt

        markdown = self._read_txt(file_path)
        parser = MarkdownIt("commonmark", {"breaks": True})
        html = parser.render(markdown)
        return self._html_to_text(html)

    def _html_to_text(self, html_input) -> str:
        from bs4 import BeautifulSoup

        if hasattr(html_input, "get_text"):
            root = html_input
        else:
            root = BeautifulSoup(str(html_input), "lxml")

        blocks: list[str] = []
        for element in root.find_all(
            ["h1", "h2", "h3", "h4", "h5", "h6", "p", "li", "pre", "code", "blockquote", "td", "th"]
        ):
            text = element.get_text(" ", strip=True)
            if text:
                blocks.append(text)

        if not blocks:
            return root.get_text("\n", strip=True)
        return "\n\n".join(blocks)

    def _make_documents(self, file_path: Path, base_dir: Path, texts: list[str]) -> list[Document]:
        documents: list[Document] = []
        for text in texts:
            cleaned = self._clean_text(text)
            if cleaned:
                documents.append(
                    Document(
                        text=cleaned,
                        metadata={"source": str(file_path.relative_to(base_dir))},
                    )
                )
        return documents

    def _clean_text(self, text: str) -> str:
        normalized = text.replace("\r\n", "\n").replace("\r", "\n")
        lines = [line.strip() for line in normalized.split("\n")]
        compact = "\n".join(line for line in lines if line)
        return compact.strip()
